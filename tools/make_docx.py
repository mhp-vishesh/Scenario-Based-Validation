"""Build the boss-facing DOCX: plan diagram + phased execution checklist.

Output: Scenario-Based_Validation_POC_Plan.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DIAGRAM = ROOT / "assets" / "poc_plan_diagram.png"
OUT = ROOT / "Scenario-Based_Validation_POC_Plan.docx"

# palette
NAVY = RGBColor(0x1F, 0x2A, 0x44)
INK = RGBColor(0x2B, 0x34, 0x47)
MUTE = RGBColor(0x5B, 0x65, 0x77)
GREEN = RGBColor(0x5C, 0x8A, 0x1B)
LINE_HEX = "D6DBE3"
HEAD_FILL = "1F2A44"
ALT_FILL = "F4F6F9"

PHASE_ACCENT = {
    "0": "5B6577",
    "1": "3B5275",
    "2": "5C8A1B",
    "3": "2C6E7F",
    "4": "B5852A",
}

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(4)

sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
PAGE_W = sec.page_width - sec.left_margin - sec.right_margin


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def set_borders(table, hex_color=LINE_HEX, size=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tblPr.append(borders)


def cell_text(cell, text, *, bold=False, color=INK, size=10, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white else color
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    return p


def pad_cell(cell, top=2, bottom=2, left=6, right=6):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val * 20))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcPr.append(mar)


def rule(color=LINE_HEX, weight=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(weight))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def heading(text, accent_hex):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(accent_hex)
    return p


# ------------------------------------------------------------------ title
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(0)
r = t.add_run("Scenario-Based Validation POC")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = NAVY

s = doc.add_paragraph()
s.paragraph_format.space_after = Pt(2)
r = s.add_run("NVIDIA Cosmos on AWS  |  plan and execution checklist")
r.font.size = Pt(11)
r.font.color.rgb = MUTE
rule(weight=14, color="1F2A44")

meta = doc.add_paragraph()
r = meta.add_run("Prepared: 18 June 2026     Domain: Autonomous driving / ADAS     Infrastructure: Self-hosted Cosmos on AWS GPUs")
r.font.size = Pt(9)
r.font.color.rgb = MUTE

# ------------------------------------------------------------------ objective
heading("Objective", "1F2A44")
doc.add_paragraph(
    "Stand up the NVIDIA Cosmos world foundation models on AWS and build a closed-loop scenario-based "
    "validation harness for an ADAS function. From a small set of seed clips we generate a large set of "
    "safety-relevant edge cases, run the system under test against them, and use an automated judge to score "
    "behaviour and surface failures. The output is a dashboard and an auditable validation report aligned to "
    "SOTIF (ISO 21448) and ISO 26262 thinking."
)

# ------------------------------------------------------------------ diagram
heading("Solution at a glance", "1F2A44")
if DIAGRAM.exists():
    doc.add_picture(str(DIAGRAM), width=PAGE_W)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cap.add_run("Figure 1. Four-phase pipeline: provision, generate, validate, present. The loop feeds new failures back into the scenario matrix.")
r.italic = True
r.font.size = Pt(8.5)
r.font.color.rgb = MUTE

# ------------------------------------------------------------------ checklist
heading("Execution checklist", "1F2A44")
intro = doc.add_paragraph()
r = intro.add_run("Owners and dates are placeholders for the kickoff meeting. Start Phase 1 quota requests on day one; they have the longest lead time.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTE

PHASES = [
    ("0", "Scope and demo narrative", "Day 1", [
        ("Confirm the single ODD slice to validate (e.g. urban intersection, pedestrian crossing)", "Lead", "One slice agreed in writing"),
        ("Select the system under test (open detector or simple planner)", "ML", "Model runs on a sample clip"),
        ("Write the one headline claim the demo must prove", "Lead", "One sentence approved"),
        ("Define success metrics (scenarios generated, failures found, coverage)", "Lead", "Metric targets set"),
    ]),
    ("1", "Provision AWS and NVIDIA AI Enterprise", "Days 1 to 3", [
        ("File GPU quota increase for p5 / p4d and g6e", "Cloud", "Quota approved by AWS"),
        ("Subscribe to NVIDIA AI Enterprise in AWS Marketplace and capture the NGC key", "Cloud", "NGC login works"),
        ("Launch generation node from the NVIDIA AI Enterprise AMI", "Cloud", "nvidia-smi and GPU docker test pass"),
        ("Launch the small judge node (L40S or A100)", "Cloud", "Node reachable, GPU visible"),
        ("Create S3 bucket and attach EBS gp3 volume", "Cloud", "Read and write verified"),
        ("docker login nvcr.io and set the Hugging Face token", "ML", "Auth confirmed"),
    ]),
    ("2", "Stand up Cosmos", "Days 3 to 4", [
        ("Clone cosmos-predict2.5, cosmos-transfer2.5, cosmos-reason2, cosmos-cookbook", "ML", "Repos cloned"),
        ("Pull model checkpoints from Hugging Face", "ML", "Checkpoints downloaded"),
        ("Run one sample inference per model from the Cookbook", "ML", "Each model produces output"),
        ("Record per-model VRAM and runtime for sizing", "ML", "Sizing table filled"),
    ]),
    ("3", "Generate scenarios", "Days 4 to 7", [
        ("Ingest seed clips and run Cosmos Curator (filter, annotate, deduplicate)", "ML", "Clean seed set ready"),
        ("Author the scenario matrix config (weather, light, actors, behaviour, geometry)", "ML", "Matrix reviewed"),
        ("Generate variants with Cosmos Transfer using structured controls", "ML", "Variants render correctly"),
        ("Roll futures with Cosmos Predict for hazardous moments", "ML", "Future clips generated"),
        ("Write every clip and its metadata to the S3 manifest", "ML", "Manifest complete and reproducible"),
    ]),
    ("4", "Validate", "Days 7 to 10", [
        ("Run the system under test on every clip and capture outputs", "ML", "Outputs stored per clip"),
        ("Configure the Cosmos Reason rubric and return structured JSON verdicts", "ML", "Verdicts parse cleanly"),
        ("Score clip realism with Cosmos Evaluator", "ML", "Realism scores attached"),
        ("Calibrate the judge against Alpamayo-style references (optional)", "ML", "Calibration documented"),
        ("Aggregate results into a single scored results table", "ML", "Results table built"),
    ]),
    ("5", "Present", "Days 10 to 12", [
        ("Build the Streamlit dashboard (metrics, coverage heatmap, failure gallery)", "ML", "Dashboard runs locally"),
        ("Implement a mock mode so the dashboard runs without a GPU", "ML", "Runs on laptop"),
        ("Add the SOTIF / ISO 26262 one-page PDF export", "ML", "PDF exports correctly"),
        ("Pre-generate all demo assets", "ML", "Demo runs offline"),
    ]),
    ("6", "Rehearse and present", "Day 12", [
        ("Dry-run the demo end to end", "Lead", "No live GPU dependency"),
        ("Lock the headline claim and the single dramatic failure clip", "Lead", "Story rehearsed"),
        ("Prepare the cost summary and the next-step ask", "Lead", "One-slide ask ready"),
    ]),
]

for num, title, dur, rows in PHASES:
    accent = PHASE_ACCENT.get(num, "1F2A44")
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(9)
    h.paragraph_format.space_after = Pt(3)
    rr = h.add_run(f"Phase {num} \u2014 {title}")
    # avoid em dash per house style: replace with hyphen
    rr.text = f"Phase {num} - {title}"
    rr.bold = True
    rr.font.size = Pt(12)
    rr.font.color.rgb = RGBColor.from_string(accent)
    rd = h.add_run(f"      {dur}")
    rd.font.size = Pt(9.5)
    rd.italic = True
    rd.font.color.rgb = MUTE

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    widths = [Inches(0.35), Inches(4.1), Inches(0.7), Inches(1.85)]
    headers = ["", "Task", "Owner", "Acceptance criteria"]
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], accent)
        cell_text(hdr[i], htext, bold=True, white=True, size=9.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else None)
        pad_cell(hdr[i])
    for ri, (task, owner, acc) in enumerate(rows):
        cells = table.add_row().cells
        if ri % 2 == 1:
            for c in cells:
                shade(c, ALT_FILL)
        cell_text(cells[0], "\u2610", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[1], task, size=9.5)
        cell_text(cells[2], owner, size=9.5, color=MUTE, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[3], acc, size=9, color=MUTE)
        for c in cells:
            pad_cell(c)
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]

# ------------------------------------------------------------------ what wow looks like
heading("What a strong result looks like", "5C8A1B")
for line in [
    "From a handful of seed clips, generate several hundred safety-relevant scenarios with no extra data collection.",
    "Surface failure modes the test fleet never encountered, each shown as a clip with the model output and a written verdict.",
    "Show coverage of the ODD as a heatmap, so gaps are visible at a glance.",
    "Export a one-page validation report that a safety lead can read and sign.",
]:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    rr = p.add_run("\u2713  ")
    rr.font.color.rgb = GREEN
    rr.bold = True
    p.add_run(line)

# ------------------------------------------------------------------ risks
heading("Key risks and mitigations", "1F2A44")
risk_tbl = doc.add_table(rows=1, cols=2)
risk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_borders(risk_tbl)
rhdr = risk_tbl.rows[0].cells
for i, htext in enumerate(["Risk", "Mitigation"]):
    shade(rhdr[i], HEAD_FILL)
    cell_text(rhdr[i], htext, bold=True, white=True, size=9.5)
    pad_cell(rhdr[i])
RISKS = [
    ("GPU quota approval is slow", "File on day one. Validate the judge node and Cosmos Reason while the large-instance quota is pending."),
    ("Generation throughput is limited", "Keep clip counts modest. Generate in short batches, snapshot to S3, stop the large node between runs."),
    ("Judge reliability is questioned", "Use a fixed rubric, add Cosmos Evaluator realism scores, and calibrate against Alpamayo-style references."),
    ("Scope creep", "One ODD slice and one system under test for the POC. Depth over breadth."),
    ("Cost overrun", "Run generation in bursts and shut the p5 / p4d node down. Keep only the small judge node warm."),
]
for ri, (risk, mit) in enumerate(RISKS):
    cells = risk_tbl.add_row().cells
    if ri % 2 == 1:
        for c in cells:
            shade(c, ALT_FILL)
    cell_text(cells[0], risk, size=9.5, bold=True)
    cell_text(cells[1], mit, size=9.5, color=MUTE)
    for c in cells:
        pad_cell(c)
for row in risk_tbl.rows:
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(4.7)

# ------------------------------------------------------------------ footer
ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = ftr.add_run("Scenario-Based Validation POC  |  NVIDIA Cosmos on AWS  |  Confidential")
fr.font.size = Pt(8)
fr.font.color.rgb = MUTE

doc.save(OUT)
print("wrote", OUT)
